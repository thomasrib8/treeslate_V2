import docx
from datetime import timedelta

def get_docx_stats(file_path):
    """Récupère les statistiques d'un fichier .docx"""
    doc = docx.Document(file_path)
    words = 0
    characters = 0
    paragraphs = len(doc.paragraphs)
    for paragraph in doc.paragraphs:
        words += len(paragraph.text.split())
        characters += len(paragraph.text)
    pages = len(doc.sections)  # Approximating pages as sections
    return words, characters, pages, paragraphs

def calculate_translation_time(words, paragraphs, group_size):
    """Calcule le temps de traduction estimé"""
    step1_time = words * 0.00156  # Temps pour l'étape 1
    group_time_map = {
        1: 1.5, 2: 1.8, 3: 2.5, 4: 3.0, 5: 3.2,
        6: 3.3, 7: 3.5, 8: 4.0, 9: 3.8, 10: 4.0
    }
    step2_time = paragraphs * group_time_map.get(group_size, 3.0)  # Par défaut 3.0 si group_size invalide
    total_time_sec = step1_time + step2_time
    return timedelta(seconds=total_time_sec), total_time_sec

def calculate_translation_cost(words, characters, translation_time_min):
    """Calcule le coût de traduction estimé"""
    tokens = words * 2
    step1_cost = tokens * 0.0000015  # Coût des tokens
    step2_cost = characters * 0.000021  # Coût DeepL
    step3_cost = translation_time_min * 0.005161  # Coût application web
    total_translation_cost = step1_cost + step2_cost + step3_cost
    return round(total_translation_cost, 6)

def calculate_review_cost(words, reviewer_choice):
    """Calcule le coût de relecture en fonction du choix du relecteur"""
    if reviewer_choice == "TOBY":
        return round(words * 0.025, 6)
    elif reviewer_choice == "TOBY+MIKE":
        return round(words * 0.025, 6)
    elif reviewer_choice == "MIKE":
        return 0.0
    else:
        raise ValueError("Invalid reviewer choice")

def get_reviewer_choice():
    """Demande à l'utilisateur de fournir un choix valide pour le relecteur"""
    valid_choices = ["TOBY", "TOBY+MIKE", "MIKE"]
    while True:
        choice = input("Enter the reviewer choice (TOBY, TOBY+MIKE, MIKE): ").strip().upper()
        if choice in valid_choices:
            return choice
        print(f"Invalid choice. Please choose from {valid_choices}.")

def main():
    # Chemin du fichier .docx
    file_path = input("Enter the path to the .docx file: ")
    words, characters, pages, paragraphs = get_docx_stats(file_path)
    
    # Affiche les statistiques du document
    print(f"Nombre de mots: {words}")
    print(f"Nombre de caractères: {characters}")
    print(f"Nombre de paragraphes: {paragraphs}")
    
    # Taille du groupe pour la traduction
    group_size = int(input("Entrez la taille du groupe (1-10): "))
    translation_time, translation_time_sec = calculate_translation_time(words, paragraphs, group_size)
    translation_time_min = translation_time_sec / 60
    
    # Calcule le coût de traduction
    translation_cost = calculate_translation_cost(words, characters, translation_time_min)
    print(f"Coût estimé de traduction: ${translation_cost:.6f}")
    
    # Demande le choix du relecteur
    reviewer_choice = get_reviewer_choice()
    review_cost = calculate_review_cost(words, reviewer_choice)
    print(f"Coût estimé de relecture: ${review_cost:.6f}")
    
    # Calcule et affiche le coût total (après arrondi des deux parties)
    total_cost = round(translation_cost + review_cost, 6)
    print(f"Total estimé du coût de traduction: ${total_cost:.6f}")

if __name__ == "__main__":
    main()
