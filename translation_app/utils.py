import requests
from docx import Document
from tqdm import tqdm
import time
import os
import pandas as pd
import logging
import openai

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def create_glossary(api_key, name, source_lang, target_lang, glossary_path):
    api_url = "https://api.deepl.com/v2/glossaries"
    
    if not os.path.exists(glossary_path):
        logger.error(f"Glossary file not found: {glossary_path}")
        raise FileNotFoundError(f"Glossary file not found: {glossary_path}")

    with open(glossary_path, "r") as glossary_file:
        glossary_content = glossary_file.read()

    headers = {
        "Authorization": f"DeepL-Auth-Key {api_key}",
        "Content-Type": "application/x-www-form-urlencoded",
    }

    data = {
        "name": name,
        "source_lang": source_lang,
        "target_lang": target_lang,
        "entries_format": "csv",
        "entries": glossary_content,
    }

    response = requests.post(api_url, headers=headers, data=data)
    response_data = response.json()

    if response.status_code in (200, 201) and "glossary_id" in response_data:
        logger.info(f"Glossary created successfully with ID: {response_data['glossary_id']}")
        return response_data["glossary_id"]
    else:
        logger.error(f"Failed to create glossary: {response.text}")
        raise Exception(f"Failed to create glossary: {response.text}")

def translate_docx_with_deepl(api_key, input_file_path, output_file_path, target_language, source_language, glossary_id=None):
    api_url = "https://api.deepl.com/v2/document"
    headers = {"Authorization": f"DeepL-Auth-Key {api_key}"}
    data = {"target_lang": target_language, "source_lang": source_language}

    if glossary_id:
        data["glossary_id"] = glossary_id

    if not os.path.exists(input_file_path):
        logger.error(f"Input document not found: {input_file_path}")
        raise FileNotFoundError(f"Input document not found: {input_file_path}")

    with open(input_file_path, "rb") as file:
        upload_response = requests.post(api_url, headers=headers, data=data, files={"file": file})

    if upload_response.status_code != 200:
        raise Exception(f"Failed to upload document: {upload_response.text}")

    upload_data = upload_response.json()
    document_id = upload_data["document_id"]
    document_key = upload_data["document_key"]

    status_url = f"{api_url}/{document_id}"
    while True:
        status_response = requests.post(status_url, headers=headers, data={"document_key": document_key})
        status_data = status_response.json()
        if status_data["status"] == "done":
            break
        elif status_data["status"] == "error":
            raise Exception(f"Translation error: {status_data}")
        time.sleep(1)

    download_url = f"{api_url}/{document_id}/result"
    download_response = requests.post(download_url, headers=headers, data={"document_key": document_key})

    if download_response.status_code == 200:
        with open(output_file_path, "wb") as output_file:
            output_file.write(download_response.content)
        logger.info(f"Translated document saved to {output_file_path}")
    else:
        raise Exception(f"Failed to download translated document: {download_response.text}")

def improve_translation(input_file, glossary_path, output_file, language_level, source_language, target_language, group_size, model):
    """
    Améliore la traduction avec ChatGPT en utilisant le glossaire.
    """
    if glossary_path and not os.path.exists(glossary_path):
        logger.error(f"Glossary file not found: {glossary_path}")
        raise FileNotFoundError(f"Glossary file not found: {glossary_path}")

    doc = Document(input_file)
    glossary = read_glossary(glossary_path) if glossary_path else {}
    output_doc = Document()
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    logger.debug(f"Loaded {len(paragraphs)} paragraphs for processing.")
    
    with tqdm(total=len(paragraphs), desc="Processing paragraphs") as pbar:
        for i in range(0, len(paragraphs), group_size):
            group = paragraphs[i : i + group_size]
            improved_text = process_paragraphs(group, glossary, language_level, source_language, target_language, model)
            if improved_text:
                output_doc.add_paragraph(improved_text)
            else:
                logger.warning(f"Skipping group {i // group_size + 1} due to an error.")
            pbar.update(len(group))
    
    output_doc.save(output_file)
    logger.debug(f"Improved document saved to {output_file}.")

def convert_excel_to_csv(excel_path, csv_path):
    if not os.path.exists(excel_path):
        logger.error(f"Excel file not found: {excel_path}")
        raise FileNotFoundError(f"Excel file not found: {excel_path}")

    df = pd.read_excel(excel_path, header=None)

    try:
        df.to_csv(csv_path, index=False, header=False, encoding='utf-8-sig')
        logger.info(f"Converted Excel file to CSV with UTF-8 encoding: {csv_path}")
    except Exception as e:
        logger.error(f"Error converting Excel to CSV: {e}")
        raise

    return csv_path


def read_glossary(glossary_path):
    """
    Lit un glossaire à partir d'un fichier Word (.docx) ou CSV.
    """
    glossary = {}
    try:
        if glossary_path.endswith(".csv"):
            df = pd.read_csv(glossary_path, header=None)
            for index, row in df.iterrows():
                glossary[row[0].strip()] = row[1].strip()
        elif glossary_path.endswith(".docx"):
            doc = Document(glossary_path)
            for paragraph in doc.paragraphs:
                if ":" in paragraph.text:
                    source, target = paragraph.text.split(":", 1)
                    glossary[source.strip()] = target.strip()
        logger.debug(f"Glossary loaded successfully from {glossary_path}.")
    except Exception as e:
        logger.error(f"Error reading glossary: {e}")
        raise
    return glossary

def process_paragraphs(paragraphs, glossary, language_level, source_language, target_language, model):
    """
    Envoie les paragraphes à ChatGPT pour amélioration de la traduction.
    """
    logger.debug(f"Processing paragraphs with model {model}.")
    prompt = (
        f"Translate the following text from {source_language} to {target_language} "
        f"and improve its quality to match the '{language_level}' language level.\n"
        f"Use the glossary strictly when applicable: {glossary}.\n"
        f"Return only the improved translation, without additional comments.\n\n"
    )

    for para in paragraphs:
        prompt += f"{para}\n\n"

    try:
        response = openai.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a skilled translator and editor."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=2048,
            temperature=0.7,
        )
        return response["choices"][0]["message"]["content"].strip()
    except openai.error.RateLimitError as e:
        logger.error(f"Rate limit reached: {e}. Adding delay before retrying.")
        time.sleep(30)
        return process_paragraphs(paragraphs, glossary, language_level, source_language, target_language, model)
    except Exception as e:
        logger.error(f"An error occurred with OpenAI API: {e}")
        raise

def ensure_directory_exists(path):
    """
    Vérifie que le répertoire du chemin donné existe, sinon le crée.
    """
    directory = os.path.dirname(path)
    if not os.path.exists(directory):
        os.makedirs(directory)
        logger.debug(f"Directory created: {directory}")
